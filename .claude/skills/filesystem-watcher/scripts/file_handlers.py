#!/usr/bin/env python3
"""
File Handlers - Metadata extraction for various file types.

Each handler extracts type-specific metadata from files.
Handlers gracefully degrade when optional dependencies are missing.
"""

from pathlib import Path
from typing import Any
import json
import struct


def extract_image_metadata(path: Path) -> dict:
    """
    Extract metadata from image files.
    
    Extracts: dimensions, format, color mode, EXIF data
    Dependencies: Pillow (optional), exifread (optional)
    """
    metadata = {}
    
    # Try Pillow for basic image info
    try:
        from PIL import Image
        from PIL.ExifTags import TAGS
        
        with Image.open(path) as img:
            metadata["dimensions"] = img.size
            metadata["format"] = img.format
            metadata["mode"] = img.mode
            
            # Extract EXIF
            exif_data = {}
            if hasattr(img, '_getexif') and img._getexif():
                for tag_id, value in img._getexif().items():
                    tag = TAGS.get(tag_id, tag_id)
                    if isinstance(value, bytes):
                        try:
                            value = value.decode('utf-8', errors='ignore')
                        except:
                            value = str(value)[:100]
                    exif_data[str(tag)] = str(value)[:200]  # Limit value length
            
            if exif_data:
                metadata["exif_data"] = exif_data
                
                # Extract common EXIF fields
                if "DateTimeOriginal" in exif_data:
                    metadata["date_taken"] = exif_data["DateTimeOriginal"]
                if "Make" in exif_data:
                    metadata["camera_make"] = exif_data["Make"]
                if "Model" in exif_data:
                    metadata["camera_model"] = exif_data["Model"]
                    
    except ImportError:
        # Fallback: try to read dimensions from file header
        metadata.update(_read_image_dimensions_fallback(path))
    except Exception as e:
        metadata["extraction_error"] = str(e)
    
    return metadata


def _read_image_dimensions_fallback(path: Path) -> dict:
    """Read image dimensions without Pillow."""
    metadata = {}
    ext = path.suffix.lower()
    
    try:
        with open(path, 'rb') as f:
            if ext in ['.jpg', '.jpeg']:
                # JPEG
                f.seek(0)
                data = f.read(30000)
                for i in range(len(data) - 8):
                    if data[i:i+2] == b'\xff\xc0':
                        height = struct.unpack('>H', data[i+5:i+7])[0]
                        width = struct.unpack('>H', data[i+7:i+9])[0]
                        metadata["dimensions"] = (width, height)
                        break
                        
            elif ext == '.png':
                # PNG
                f.seek(16)
                width = struct.unpack('>I', f.read(4))[0]
                height = struct.unpack('>I', f.read(4))[0]
                metadata["dimensions"] = (width, height)
                
            elif ext == '.gif':
                # GIF
                f.seek(6)
                width = struct.unpack('<H', f.read(2))[0]
                height = struct.unpack('<H', f.read(2))[0]
                metadata["dimensions"] = (width, height)
    except:
        pass
    
    return metadata


def extract_pdf_metadata(path: Path) -> dict:
    """
    Extract metadata from PDF files.
    
    Extracts: page count, title, author, text preview
    Dependencies: pypdf (optional), pdfplumber (optional)
    """
    metadata = {}
    
    # Try pypdf
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(str(path))
        metadata["page_count"] = len(reader.pages)
        
        # Document info
        if reader.metadata:
            if reader.metadata.title:
                metadata["title"] = reader.metadata.title
            if reader.metadata.author:
                metadata["author"] = reader.metadata.author
            if reader.metadata.subject:
                metadata["subject"] = reader.metadata.subject
            if reader.metadata.creator:
                metadata["creator"] = reader.metadata.creator
        
        # Extract text preview from first page
        try:
            first_page = reader.pages[0]
            text = first_page.extract_text() or ""
            metadata["content_preview"] = text[:1000].strip()
        except:
            pass
            
        return metadata
        
    except ImportError:
        pass
    
    # Try pdfplumber as fallback
    try:
        import pdfplumber
        
        with pdfplumber.open(str(path)) as pdf:
            metadata["page_count"] = len(pdf.pages)
            
            if pdf.metadata:
                metadata.update({
                    k: str(v)[:200] for k, v in pdf.metadata.items() 
                    if v and k in ['Title', 'Author', 'Subject', 'Creator']
                })
            
            # Text preview
            try:
                first_page = pdf.pages[0]
                text = first_page.extract_text() or ""
                metadata["content_preview"] = text[:1000].strip()
            except:
                pass
                
        return metadata
        
    except ImportError:
        metadata["note"] = "Install pypdf or pdfplumber for PDF metadata extraction"
    except Exception as e:
        metadata["extraction_error"] = str(e)
    
    return metadata


def extract_document_metadata(path: Path) -> dict:
    """
    Extract metadata from document files.
    
    Handles: .docx, .txt, .md, .rtf
    Dependencies: python-docx (optional)
    """
    metadata = {}
    ext = path.suffix.lower()
    
    if ext == '.pdf':
        return extract_pdf_metadata(path)
    
    if ext in ['.txt', '.md']:
        try:
            content = path.read_text(encoding='utf-8', errors='ignore')
            metadata["content_preview"] = content[:1000].strip()
            metadata["line_count"] = content.count('\n') + 1
            metadata["word_count"] = len(content.split())
            metadata["char_count"] = len(content)
        except Exception as e:
            metadata["extraction_error"] = str(e)
        return metadata
    
    if ext == '.docx':
        try:
            from docx import Document
            
            doc = Document(str(path))
            
            # Core properties
            props = doc.core_properties
            if props.title:
                metadata["title"] = props.title
            if props.author:
                metadata["author"] = props.author
            if props.subject:
                metadata["subject"] = props.subject
            if props.created:
                metadata["doc_created"] = props.created.isoformat()
            if props.modified:
                metadata["doc_modified"] = props.modified.isoformat()
            
            # Extract text preview
            paragraphs = []
            for para in doc.paragraphs[:20]:  # First 20 paragraphs
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            content = "\n".join(paragraphs)
            metadata["content_preview"] = content[:1000]
            metadata["paragraph_count"] = len(doc.paragraphs)
            
            return metadata
            
        except ImportError:
            metadata["note"] = "Install python-docx for DOCX metadata extraction"
        except Exception as e:
            metadata["extraction_error"] = str(e)
    
    return metadata


def extract_audio_metadata(path: Path) -> dict:
    """
    Extract metadata from audio files.
    
    Extracts: duration, bitrate, sample rate, artist, title, album
    Dependencies: mutagen (optional), tinytag (optional)
    """
    metadata = {}
    
    # Try mutagen
    try:
        from mutagen import File
        from mutagen.mp3 import MP3
        from mutagen.flac import FLAC
        
        audio = File(str(path))
        
        if audio:
            metadata["duration"] = audio.info.length if hasattr(audio.info, 'length') else None
            
            if hasattr(audio.info, 'bitrate'):
                metadata["bitrate"] = audio.info.bitrate
            if hasattr(audio.info, 'sample_rate'):
                metadata["sample_rate"] = audio.info.sample_rate
            if hasattr(audio.info, 'channels'):
                metadata["channels"] = audio.info.channels
            
            # Tags
            if hasattr(audio, 'tags') and audio.tags:
                tags = {}
                for key in ['title', 'artist', 'album', 'date', 'genre', 'tracknumber']:
                    if key in audio.tags:
                        tags[key] = str(audio.tags[key][0])[:200]
                    elif key.upper() in audio.tags:
                        tags[key] = str(audio.tags[key.upper()][0])[:200]
                metadata.update(tags)
                
        return metadata
        
    except ImportError:
        pass
    
    # Try tinytag as fallback
    try:
        from tinytag import TinyTag
        
        tag = TinyTag.get(str(path))
        
        metadata["duration"] = tag.duration
        metadata["bitrate"] = tag.bitrate
        metadata["sample_rate"] = tag.samplerate
        metadata["channels"] = tag.channels
        
        if tag.title:
            metadata["title"] = tag.title
        if tag.artist:
            metadata["artist"] = tag.artist
        if tag.album:
            metadata["album"] = tag.album
        if tag.year:
            metadata["year"] = tag.year
        if tag.genre:
            metadata["genre"] = tag.genre
            
        return metadata
        
    except ImportError:
        metadata["note"] = "Install mutagen or tinytag for audio metadata extraction"
    except Exception as e:
        metadata["extraction_error"] = str(e)
    
    return metadata


def extract_video_metadata(path: Path) -> dict:
    """
    Extract metadata from video files.
    
    Extracts: duration, dimensions, codec, fps
    Dependencies: ffprobe (optional), moviepy (optional)
    """
    metadata = {}
    
    # Try ffprobe via subprocess
    try:
        import subprocess
        import json
        
        cmd = [
            'ffprobe',
            '-v', 'quiet',
            '-print_format', 'json',
            '-show_format',
            '-show_streams',
            str(path)
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0:
            data = json.loads(result.stdout)
            
            # Format info
            if 'format' in data:
                fmt = data['format']
                if 'duration' in fmt:
                    metadata["duration"] = float(fmt['duration'])
                if 'bit_rate' in fmt:
                    metadata["bitrate"] = int(fmt['bit_rate'])
            
            # Stream info
            for stream in data.get('streams', []):
                if stream.get('codec_type') == 'video':
                    metadata["dimensions"] = (
                        stream.get('width'),
                        stream.get('height')
                    )
                    metadata["codec"] = stream.get('codec_name')
                    
                    # Calculate FPS
                    if 'r_frame_rate' in stream:
                        num, den = map(int, stream['r_frame_rate'].split('/'))
                        if den > 0:
                            metadata["fps"] = round(num / den, 2)
                    break
                    
            return metadata
            
    except (ImportError, FileNotFoundError, subprocess.TimeoutExpired):
        pass
    except Exception as e:
        metadata["extraction_error"] = str(e)
    
    # Try moviepy as fallback
    try:
        from moviepy.editor import VideoFileClip
        
        with VideoFileClip(str(path)) as clip:
            metadata["duration"] = clip.duration
            metadata["dimensions"] = clip.size
            metadata["fps"] = clip.fps
            
        return metadata
        
    except ImportError:
        metadata["note"] = "Install ffprobe or moviepy for video metadata extraction"
    except Exception as e:
        metadata["extraction_error"] = str(e)
    
    return metadata


def extract_archive_metadata(path: Path) -> dict:
    """
    Extract metadata from archive files.
    
    Extracts: file count, total size, file listing
    Dependencies: zipfile (builtin), tarfile (builtin)
    """
    metadata = {}
    ext = path.suffix.lower()
    
    try:
        if ext == '.zip':
            import zipfile
            with zipfile.ZipFile(str(path), 'r') as zf:
                files = zf.namelist()
                metadata["file_count"] = len(files)
                metadata["total_uncompressed"] = sum(
                    info.file_size for info in zf.infolist()
                )
                metadata["file_list"] = files[:50]  # First 50 files
                
        elif ext in ['.tar', '.gz', '.bz2', '.xz']:
            import tarfile
            with tarfile.open(str(path), 'r:*') as tf:
                members = tf.getmembers()
                metadata["file_count"] = len(members)
                metadata["total_uncompressed"] = sum(
                    m.size for m in members if m.isfile()
                )
                metadata["file_list"] = [m.name for m in members[:50]]
                
        elif ext == '.7z':
            try:
                import py7zr
                with py7zr.SevenZipFile(str(path), 'r') as zf:
                    files = zf.getnames()
                    metadata["file_count"] = len(files)
                    metadata["file_list"] = files[:50]
            except ImportError:
                metadata["note"] = "Install py7zr for 7z metadata extraction"
                
        elif ext == '.rar':
            try:
                import rarfile
                with rarfile.RarFile(str(path), 'r') as rf:
                    files = rf.namelist()
                    metadata["file_count"] = len(files)
                    metadata["file_list"] = files[:50]
            except ImportError:
                metadata["note"] = "Install rarfile for RAR metadata extraction"
                
    except Exception as e:
        metadata["extraction_error"] = str(e)
    
    return metadata


def extract_code_metadata(path: Path) -> dict:
    """
    Extract metadata from code files.
    
    Extracts: line count, imports, functions/classes
    """
    metadata = {}
    
    try:
        content = path.read_text(encoding='utf-8', errors='ignore')
        lines = content.split('\n')
        
        metadata["line_count"] = len(lines)
        metadata["content_preview"] = content[:1000].strip()
        
        ext = path.suffix.lower()
        
        if ext == '.py':
            # Python-specific
            imports = [l.strip() for l in lines if l.strip().startswith(('import ', 'from '))]
            functions = [l.strip() for l in lines if l.strip().startswith('def ')]
            classes = [l.strip() for l in lines if l.strip().startswith('class ')]
            
            metadata["imports"] = imports[:20]
            metadata["function_count"] = len(functions)
            metadata["class_count"] = len(classes)
            
        elif ext in ['.js', '.ts']:
            # JavaScript/TypeScript
            imports = [l.strip() for l in lines if 'import ' in l or 'require(' in l]
            functions = [l.strip() for l in lines if 'function ' in l or '=>' in l]
            
            metadata["imports"] = imports[:20]
            metadata["function_count"] = len(functions)
            
    except Exception as e:
        metadata["extraction_error"] = str(e)
    
    return metadata


def extract_data_metadata(path: Path) -> dict:
    """
    Extract metadata from data files.
    
    Handles: JSON, YAML, CSV, XML
    """
    metadata = {}
    ext = path.suffix.lower()
    
    try:
        if ext == '.json':
            content = path.read_text(encoding='utf-8')
            data = json.loads(content)
            
            if isinstance(data, list):
                metadata["type"] = "array"
                metadata["item_count"] = len(data)
            elif isinstance(data, dict):
                metadata["type"] = "object"
                metadata["key_count"] = len(data)
                metadata["keys"] = list(data.keys())[:20]
            
            metadata["content_preview"] = content[:1000]
            
        elif ext in ['.yaml', '.yml']:
            try:
                import yaml
                content = path.read_text(encoding='utf-8')
                data = yaml.safe_load(content)
                
                if isinstance(data, list):
                    metadata["type"] = "array"
                    metadata["item_count"] = len(data)
                elif isinstance(data, dict):
                    metadata["type"] = "object"
                    metadata["key_count"] = len(data)
                    metadata["keys"] = list(data.keys())[:20]
                    
                metadata["content_preview"] = content[:1000]
            except ImportError:
                metadata["note"] = "Install pyyaml for YAML metadata extraction"
                
        elif ext == '.csv':
            import csv
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f)
                headers = next(reader, [])
                row_count = sum(1 for _ in reader) + 1
                
            metadata["headers"] = headers
            metadata["column_count"] = len(headers)
            metadata["row_count"] = row_count
            
    except Exception as e:
        metadata["extraction_error"] = str(e)
    
    return metadata
